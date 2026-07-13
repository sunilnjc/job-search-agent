from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

# Sections that don't change per job — mirrors resumes/Sunilkumar Kalabandi FDE v2.docx.
# If that base resume is edited, update these to match.

HEADER_NAME = "SUNILKUMAR KALABANDI"
HEADER_CONTACT = (
    "Dubai, UAE  •  +971-585077967  •  sunilkumar.kalabandi@gmail.com  •  "
    "linkedin.com/in/sunilkumar-kalabandi  •  github.com/sunilnjc"
)

# The header tagline was previously a hardcoded "Forward Deployed Engineer..." line
# applied to every resume regardless of the job — wrong and confusing for a generic
# Senior Software Engineer role. Pick a tagline based on the target job's title instead.
TAGLINE_BY_CATEGORY = {
    "fde": "Forward Deployed Engineer  |  Solutions Architect  |  Payments & Banking Platforms",
    "senior": "Senior Software Engineer  |  Distributed Systems  |  Banking Platforms  |  Cloud Migration",
    "software": "Software Engineer  |  Distributed Systems  |  Banking Platforms",
}


def _categorize_job_title(title: str) -> str:
    t = title.lower()
    if "forward deployed" in t:
        return "fde"
    if "senior" in t or "staff" in t or "sr." in t or re.search(r"\bsr\b", t):
        return "senior"
    return "software"


def header_tagline(job_title: str) -> str:
    return TAGLINE_BY_CATEGORY[_categorize_job_title(job_title)]

CORE_COMPETENCIES = [
    "Domain: Payment systems, banking/fintech, regulatory compliance (UAE/KSA/UK/Singapore), "
    "cross-border transactions, settlement, idempotency",
    "Technical: Java, Spring Boot, Kafka, Kubernetes/OpenShift, microservices, distributed systems, "
    "REST API design, MongoDB, PostgreSQL, Redis, Python, Node.js",
    "AI & LLM tooling: LLM API integration (OpenAI, Anthropic, local models via Ollama), embeddings "
    "& semantic search, prompt pipelines, agent workflows",
    "Customer-facing: requirements translation, integration architecture, partner onboarding & "
    "enablement, stakeholder communication",
]

EXPERIENCE = [
    {
        "role": "Senior Software Engineer  |  Emirates NBD (via Bitech Middle East), Dubai",
        "dates": "Oct 2023 – Present",
        "bullets": [
            "Technical lead for KSAX Localization: embedded with infrastructure, security, compliance, "
            "and platform teams to migrate 20+ mission-critical banking services to OCI KSA with a "
            "zero-disruption cutover.",
            "Designed reusable integration and automation frameworks (Vault migration, Kafka ACL "
            "provisioning, deployment orchestration) adopted by parallel teams, reducing manual effort "
            "by ~90%.",
            "Delivered Global Transactions, Digital Signatures, and Push Notifications services; shaped "
            "architecture decisions across banking domains and unblocked partner integrations.",
        ],
        "tech": "Technologies: Java, Spring Boot, Kafka, Redis, MongoDB, Node.js, OCI, OpenShift",
    },
    {
        "role": "Senior Java Developer  |  ADSS (via Halian), Abu Dhabi",
        "dates": "Nov 2022 – Sep 2023",
        "bullets": [
            "Designed enterprise reporting and data-warehousing solutions supporting business analytics "
            "and operational visibility for a regulated trading firm.",
            "Built ETL pipelines with Databricks and Apache Spark, improving reporting performance and "
            "data accessibility; partnered with business stakeholders to turn reporting requirements "
            "into scalable solutions.",
        ],
        "tech": "Technologies: Java, Spring Boot, Databricks, Apache Spark, SQL",
    },
    {
        "role": "Software Engineer  |  Emirates NBD (via Bitech Middle East), Dubai",
        "dates": "Jul 2021 – Nov 2022",
        "bullets": [
            "Owned production reliability and user experience for E20 Banking and Tablet Banking 2.0 "
            "platforms, shipping 40+ enhancements and resolving critical production issues.",
            "Built operational dashboards and monitoring (ELK, Grafana) enabling proactive issue "
            "detection, cutting mean-time-to-resolution from hours to minutes and improving uptime to "
            "99.99%+.",
        ],
        "tech": "Technologies: Java, Spring Boot, ELK, Grafana, OpenShift",
    },
    {
        "role": "Software Engineer  |  Arab Bank (via 3i Infotech), Dubai",
        "dates": "Jan 2021 – Jul 2021",
        "bullets": [
            "Core engineer on the Reflect digital bank: 100+ microservice platform launched in 4 months. "
            "Designed APIs for onboarding, wallet, credit card issuance, and payments; coordinated with "
            "operations and compliance to launch.",
            "Built scalable REST APIs and integration documentation that let partner banks go live in "
            "days rather than months.",
        ],
        "tech": "Technologies: Java, Spring Boot, Kafka, MongoDB",
    },
    {
        "role": "Software Engineer  |  Emirates NBD (via 3i Infotech), Dubai",
        "dates": "Jul 2019 – Jan 2021",
        "bullets": [
            "Designed and built Beneficiaries & Transfers microservices powering digital banking for "
            "millions of users across UAE, KSA, London, and Singapore.",
            "Worked directly with partner banks, fintechs, and payment networks on integrations — "
            "translating their requirements into API design and handling FX, compliance, and settlement "
            "edge cases.",
        ],
        "tech": None,
    },
]

EARLIER_EXPERIENCE = [
    "Associate Application Developer, Fujitsu Consulting, India (Sep 2017 – Jul 2019): Japan Next "
    "Generation Banking initiatives and enterprise inventory and operations applications.",
    "Software Engineer, Capgemini, India (Sep 2014 – May 2017): Enterprise financial services "
    "applications, payment processing, and credit card integrations for Synchrony Financial and Fiserv.",
]

PROJECTS = [
    "Job Search Agent (2026): Full-stack agent that aggregates postings from public job APIs, ranks fit "
    "against a parsed resume using local LLM embeddings (Ollama, nomic-embed-text) with an LLM scoring "
    "pass, and generates tailored application drafts (cover letters + resumes) via OpenAI/Anthropic APIs. "
    "Python CLI + FastAPI backend over a SQLite pipeline, with a React/TypeScript kanban web UI for "
    "reviewing matches, triggering fetch/match runs, and tracking applications through each stage. "
    "github.com/sunilnjc",
    "BudgetTracker (2026): Personal finance app for budgets, expenses, debts, and recurring payments "
    "with multi-month spending projections — built to replace subscription budgeting apps with a "
    "self-hosted, AI-assisted tool. React/TypeScript dashboard over a local API, developed with "
    "AI-assisted engineering workflows (Claude, OpenAI).",
]

EDUCATION = [
    "Bachelor of Engineering, Computer Science — Chaitanya Bharathi Institute of Technology (CBIT), "
    "Hyderabad, India. Graduated with Distinction, First Class.",
    "Microsoft Certified: Azure Fundamentals (AZ-900)",
]


def parse_tailoring_notes(markdown: str) -> tuple[str, list[str]]:
    """Extract the tailored summary and bullet points from an LLM-generated resume_tailoring.md.

    The model doesn't format this with perfect consistency run to run (### vs **, "Tailored
    Bullet Points" vs "Tailored Experience Bullet Points"), so match loosely on the meaningful
    words and ignore markdown decoration.
    """
    summary_header = re.search(r"^.*Tailored Professional Summary.*$", markdown, re.MULTILINE | re.IGNORECASE)
    bullets_header = re.search(r"^.*Tailored.*Bullet Points.*$", markdown, re.MULTILINE | re.IGNORECASE)

    if not summary_header or not bullets_header:
        return "", []

    summary = markdown[summary_header.end():bullets_header.start()].strip()
    bullets_section = markdown[bullets_header.end():]
    bullets = [
        line.lstrip("-•* ").strip()
        for line in bullets_section.splitlines()
        if line.strip().lstrip("*").strip().startswith(("-", "•"))
    ]
    return summary, bullets


def _para(doc: Document, text: str, bold: bool = False, size: float | None = None, space_after: int = 4, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def _heading(doc: Document, text: str):
    return _para(doc, text, bold=True, size=11.5, space_after=3)


def _bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    return p


def build_tailored_resume(summary: str, highlights: list[str], output_path: Path, job_title: str = "") -> None:
    """Generate a tailored resume docx: fixed sections from the v2 base resume, with the
    professional summary, career highlights, and header tagline swapped for job-specific ones."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.6)
        section.left_margin = section.right_margin = Inches(0.7)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    _para(doc, HEADER_NAME, bold=True, size=16, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, header_tagline(job_title), bold=True, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, HEADER_CONTACT, size=9.5, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    _heading(doc, "PROFESSIONAL SUMMARY")
    _para(doc, summary, space_after=8)

    _heading(doc, "CAREER HIGHLIGHTS")
    for bullet in highlights:
        _bullet(doc, bullet)
    _para(doc, "", space_after=4)

    _heading(doc, "CORE COMPETENCIES")
    for line in CORE_COMPETENCIES:
        _bullet(doc, line)
    _para(doc, "", space_after=4)

    _heading(doc, "PROFESSIONAL EXPERIENCE")
    for job in EXPERIENCE:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(job["role"])
        r.bold = True
        p.add_run(f"\t{job['dates']}")
        from docx.enum.text import WD_TAB_ALIGNMENT

        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.1), WD_TAB_ALIGNMENT.RIGHT)
        for bullet in job["bullets"]:
            _bullet(doc, bullet)
        if job["tech"]:
            _para(doc, job["tech"], size=9.5, space_after=6)
        else:
            _para(doc, "", space_after=4)

    _heading(doc, "EARLIER EXPERIENCE")
    for line in EARLIER_EXPERIENCE:
        _bullet(doc, line)
    _para(doc, "", space_after=4)

    _heading(doc, "PROJECTS")
    for line in PROJECTS:
        _bullet(doc, line)
    _para(doc, "", space_after=4)

    _heading(doc, "EDUCATION & CERTIFICATIONS")
    for line in EDUCATION:
        _bullet(doc, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def build_tailored_resume_pdf(summary: str, highlights: list[str], output_path: Path, job_title: str = "") -> None:
    """Same content as build_tailored_resume, rendered directly as PDF.

    Many application portals (especially German/EU corporate ones) reject .docx uploads
    outright or enforce it silently via a vague "check format and size" error — PDF is
    the one format almost universally accepted, so this is the one to actually attach.
    """
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    styles = {
        "name": ParagraphStyle("name", fontName="Helvetica-Bold", fontSize=16, alignment=TA_CENTER, spaceAfter=2),
        "tagline": ParagraphStyle("tagline", fontName="Helvetica-Bold", fontSize=10.5, alignment=TA_CENTER, spaceAfter=2),
        "contact": ParagraphStyle("contact", fontName="Helvetica", fontSize=8.5, alignment=TA_CENTER, spaceAfter=10),
        "heading": ParagraphStyle("heading", fontName="Helvetica-Bold", fontSize=11, spaceBefore=8, spaceAfter=4),
        "body": ParagraphStyle("body", fontName="Helvetica", fontSize=9.5, leading=13, spaceAfter=6),
        "bullet": ParagraphStyle("bullet", fontName="Helvetica", fontSize=9.5, leading=13),
        "role": ParagraphStyle("role", fontName="Helvetica-Bold", fontSize=9.5, spaceBefore=4, spaceAfter=2),
        "tech": ParagraphStyle("tech", fontName="Helvetica", fontSize=8.5, textColor="#444444", spaceAfter=6),
    }

    def bullets(lines: list[str]):
        return ListFlowable(
            [ListItem(Paragraph(line, styles["bullet"]), leftIndent=10) for line in lines],
            bulletType="bullet",
            leftIndent=14,
        )

    story = [
        Paragraph(HEADER_NAME, styles["name"]),
        Paragraph(header_tagline(job_title), styles["tagline"]),
        Paragraph(HEADER_CONTACT, styles["contact"]),
        Paragraph("PROFESSIONAL SUMMARY", styles["heading"]),
        Paragraph(summary, styles["body"]),
        Paragraph("CAREER HIGHLIGHTS", styles["heading"]),
        bullets(highlights),
        Spacer(1, 6),
        Paragraph("CORE COMPETENCIES", styles["heading"]),
        bullets(CORE_COMPETENCIES),
        Spacer(1, 6),
        Paragraph("PROFESSIONAL EXPERIENCE", styles["heading"]),
    ]
    for job in EXPERIENCE:
        story.append(Paragraph(f"{job['role']}&nbsp;&nbsp;&nbsp;&nbsp;{job['dates']}", styles["role"]))
        story.append(bullets(job["bullets"]))
        if job["tech"]:
            story.append(Paragraph(job["tech"], styles["tech"]))
        else:
            story.append(Spacer(1, 6))

    story.append(Paragraph("EARLIER EXPERIENCE", styles["heading"]))
    story.append(bullets(EARLIER_EXPERIENCE))
    story.append(Spacer(1, 6))
    story.append(Paragraph("PROJECTS", styles["heading"]))
    story.append(bullets(PROJECTS))
    story.append(Spacer(1, 6))
    story.append(Paragraph("EDUCATION & CERTIFICATIONS", styles["heading"]))
    story.append(bullets(EDUCATION))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch,
    )
    doc.build(story)
